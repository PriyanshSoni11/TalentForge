import io
import re

from pypdf import PdfReader
from docx import Document


def extract_text_from_resume(file_bytes, filename):
    if not file_bytes:
        raise ValueError("The uploaded resume file is empty.")

    lower = (filename or "").lower()

    if lower.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                except Exception:
                    continue
            full_text = "\n".join(pages_text).strip()
            # Clean null bytes and control chars
            full_text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", full_text)
            if not full_text:
                raise ValueError("Could not extract readable text from the PDF. If it is a scanned image, please upload a text-based PDF or DOCX.")
            return full_text
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"Failed to read PDF file: {exc}")

    if lower.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            full_text = "\n".join(paragraphs).strip()
            full_text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", full_text)
            if not full_text:
                raise ValueError("The DOCX file does not contain readable text.")
            return full_text
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"Failed to read DOCX file: {exc}")

    if lower.endswith(".txt"):
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode text file.")

    if lower.endswith(".doc"):
        raise ValueError("Legacy .doc format is not supported. Please save/export your resume as .pdf or .docx and upload again.")

    raise ValueError("Unsupported resume format — only .pdf and .docx are accepted.")

